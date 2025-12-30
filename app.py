# app.py — Olivetti Desk (one file, paste+click)
import os
import re
import math
import json
import time
import hashlib
from datetime import datetime
from typing import List, Dict, Any, Optional

import streamlit as st

# ============================================================
# ENV / METADATA HYGIENE
# ============================================================
os.environ.setdefault("MS_APP_ID", "olivetti-writing-desk")
os.environ.setdefault("ms-appid", "olivetti-writing-desk")

DEFAULT_MODEL = "gpt-4.1-mini"


def _get_openai_key_or_empty() -> str:
    try:
        return str(st.secrets.get("OPENAI_API_KEY", ""))  # type: ignore[attr-defined]
    except Exception:
        return ""


def _get_openai_model() -> str:
    try:
        return str(st.secrets.get("OPENAI_MODEL", DEFAULT_MODEL))  # type: ignore[attr-defined]
    except Exception:
        return os.getenv("OPENAI_MODEL", DEFAULT_MODEL)


OPENAI_MODEL = _get_openai_model()


def has_openai_key() -> bool:
    return bool(os.getenv("OPENAI_API_KEY") or _get_openai_key_or_empty())


def require_openai_key() -> str:
    key = os.getenv("OPENAI_API_KEY") or _get_openai_key_or_empty()
    if not key:
        st.error(
            "OPENAI_API_KEY is not set. Add it as an environment variable (OPENAI_API_KEY) "
            "or as a Streamlit secret."
        )
        st.stop()
    return key


# ============================================================
# CONFIG
# ============================================================
st.set_page_config(
    page_title="Olivetti Desk", layout="wide", initial_sidebar_state="expanded"
)

# Keyboard shortcut handler for undo/redo
KEYBOARD_SHORTCUTS = """
<script>
document.addEventListener('keydown', function(e) {
    // Ctrl+Z or Cmd+Z for Undo
    if ((e.ctrlKey || e.metaKey) && e.key === 'z' && !e.shiftKey) {
        e.preventDefault();
        const undoButton = window.parent.document.querySelector('[data-testid="baseButton-secondary"]:not([disabled])');
        if (undoButton && undoButton.textContent.includes('Undo')) {
            undoButton.click();
        }
    }
    // Ctrl+Y or Cmd+Shift+Z for Redo
    if (((e.ctrlKey || e.metaKey) && e.key === 'y') || ((e.ctrlKey || e.metaKey) && e.shiftKey && e.key === 'z')) {
        e.preventDefault();
        const buttons = window.parent.document.querySelectorAll('[data-testid="baseButton-secondary"]:not([disabled])');
        for (let btn of buttons) {
            if (btn.textContent.includes('Redo')) {
                btn.click();
                break;
            }
        }
    }
});
</script>
"""
st.markdown(KEYBOARD_SHORTCUTS, unsafe_allow_html=True)

# ============================================================
# GLOBALS
# ============================================================
LANES = ["Dialogue", "Narration", "Interiority", "Action"]
BAYS = ["NEW", "ROUGH", "EDIT", "FINAL"]
ENGINE_STYLES = ["NARRATIVE", "DESCRIPTIVE", "EMOTIONAL", "LYRICAL"]
AUTOSAVE_DIR = "autosave"
AUTOSAVE_PATH = os.path.join(AUTOSAVE_DIR, "olivetti_state.json")
MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10MB guardrail
WORD_RE = re.compile(r"[A-Za-z']+")
AUTOSAVE_MIN_INTERVAL_S = 12.0  # throttle rerun autosaves


# ============================================================
# SYSTEM CONTRACT - PROJECT INTEGRITY
# ============================================================
def init_system_contract() -> Dict[str, Any]:
    """Initialize the System Contract that governs all project behavior."""
    return {
        "version": "1.0",
        "created_ts": now_ts(),
        "contract_rules": {
            "user_instructions_are_law": True,
            "project_settings_override_learning": True,
            "intensity_controls_are_hard_limits": True,
            "no_cross_project_leakage": True,
            "consistency_across_sessions": True,
        },
        "project_settings": {
            "locked": False,
            "ai_intensity_max": 1.0,
            "ai_intensity_min": 0.0,
            "writing_style_locked": False,
            "voice_selection_locked": False,
            "learning_boundaries": {
                "respect_intensity_limits": True,
                "respect_style_locks": True,
                "respect_voice_locks": True,
            },
        },
        "isolation": {
            "project_id": None,
            "bay_isolation_enabled": True,
            "voice_isolation_per_bay": True,
        },
    }


def validate_contract_compliance(
    operation: str, current_settings: Dict[str, Any]
) -> bool:
    """Validate that an operation complies with the System Contract."""
    contract = st.session_state.get("_system_contract")
    if not contract:
        return True  # No contract, allow operation

    project_settings = contract.get("project_settings", {})

    # Check if project settings are locked
    if project_settings.get("locked") and operation in [
        "change_style",
        "change_voice",
        "change_intensity",
    ]:
        return False

    # Check intensity limits
    if operation == "set_intensity":
        intensity = current_settings.get("intensity", 0.0)
        max_intensity = project_settings.get("ai_intensity_max", 1.0)
        min_intensity = project_settings.get("ai_intensity_min", 0.0)
        if intensity > max_intensity or intensity < min_intensity:
            return False

    # Check style locks
    if operation == "change_style" and project_settings.get("writing_style_locked"):
        return False

    # Check voice locks
    if operation == "change_voice" and project_settings.get("voice_selection_locked"):
        return False

    return True


def enforce_contract_on_learning(learning_data: Dict[str, Any]) -> Dict[str, Any]:
    """Enforce System Contract rules on learning data before applying."""
    contract = st.session_state.get("_system_contract")
    if not contract:
        return learning_data

    boundaries = contract.get("project_settings", {}).get("learning_boundaries", {})

    # If learning must respect locks, check and constrain
    if boundaries.get("respect_intensity_limits"):
        # Ensure learned patterns don't suggest exceeding intensity limits
        max_intensity = contract.get("project_settings", {}).get(
            "ai_intensity_max", 1.0
        )
        if learning_data.get("suggested_intensity", 0.0) > max_intensity:
            learning_data["suggested_intensity"] = max_intensity

    return learning_data


def lock_project_settings() -> None:
    """Lock all project settings to prevent changes."""
    contract = st.session_state.get("_system_contract")
    if not contract:
        contract = init_system_contract()
        st.session_state._system_contract = contract

    contract["project_settings"]["locked"] = True
    mark_dirty()


def unlock_project_settings() -> None:
    """Unlock project settings to allow changes."""
    contract = st.session_state.get("_system_contract")
    if contract:
        contract["project_settings"]["locked"] = False
        mark_dirty()


def set_intensity_limits(min_val: float, max_val: float) -> None:
    """Set hard limits for AI intensity."""
    contract = st.session_state.get("_system_contract")
    if not contract:
        contract = init_system_contract()
        st.session_state._system_contract = contract

    contract["project_settings"]["ai_intensity_min"] = max(0.0, min(min_val, 1.0))
    contract["project_settings"]["ai_intensity_max"] = max(0.0, min(max_val, 1.0))

    # Enforce current intensity within new limits
    current_intensity = st.session_state.get("ai_intensity", 0.75)
    if current_intensity < contract["project_settings"]["ai_intensity_min"]:
        st.session_state.ai_intensity = contract["project_settings"]["ai_intensity_min"]
    elif current_intensity > contract["project_settings"]["ai_intensity_max"]:
        st.session_state.ai_intensity = contract["project_settings"]["ai_intensity_max"]

    mark_dirty()


def lock_writing_style() -> None:
    """Lock writing style to prevent changes."""
    contract = st.session_state.get("_system_contract")
    if not contract:
        contract = init_system_contract()
        st.session_state._system_contract = contract

    contract["project_settings"]["writing_style_locked"] = True
    mark_dirty()


def unlock_writing_style() -> None:
    """Unlock writing style to allow changes."""
    contract = st.session_state.get("_system_contract")
    if contract:
        contract["project_settings"]["writing_style_locked"] = False
        mark_dirty()


def lock_voice_selection() -> None:
    """Lock voice selection to prevent changes."""
    contract = st.session_state.get("_system_contract")
    if not contract:
        contract = init_system_contract()
        st.session_state._system_contract = contract

    contract["project_settings"]["voice_selection_locked"] = True
    mark_dirty()


def unlock_voice_selection() -> None:
    """Unlock voice selection to allow changes."""
    contract = st.session_state.get("_system_contract")
    if contract:
        contract["project_settings"]["voice_selection_locked"] = False
        mark_dirty()


def reset_system_contract() -> None:
    """Reset System Contract to defaults."""
    st.session_state._system_contract = init_system_contract()
    mark_dirty()


def get_contract_status() -> Dict[str, Any]:
    """Get current System Contract status."""
    contract = st.session_state.get("_system_contract")
    if not contract:
        return {
            "exists": False,
            "locked": False,
            "intensity_limits": {"min": 0.0, "max": 1.0},
            "style_locked": False,
            "voice_locked": False,
        }

    project_settings = contract.get("project_settings", {})
    return {
        "exists": True,
        "locked": project_settings.get("locked", False),
        "intensity_limits": {
            "min": project_settings.get("ai_intensity_min", 0.0),
            "max": project_settings.get("ai_intensity_max", 1.0),
        },
        "style_locked": project_settings.get("writing_style_locked", False),
        "voice_locked": project_settings.get("voice_selection_locked", False),
    }


# ============================================================
# UTILS
# ============================================================
def now_ts() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def now_unix() -> float:
    return time.time()


def _normalize_text(s: str) -> str:
    t = (s or "").strip()
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    t = re.sub(r"\n{3,}", "\n\n", t)
    t = re.sub(r"[ \t]{2,}", " ", t)
    return t.strip()


def _split_paragraphs(text: str) -> List[str]:
    t = _normalize_text(text)
    if not t:
        return []
    return [p.strip() for p in re.split(r"\n\s*\n", t, flags=re.MULTILINE) if p.strip()]


def _safe_filename(s: str, fallback: str = "olivetti") -> str:
    s = re.sub(r"[^\w\- ]+", "", (s or "").strip()).strip()
    s = re.sub(r"\s+", "_", s)
    return s[:80] if s else fallback


def _clamp_text(s: str, max_chars: int = 12000) -> str:
    s = s or ""
    if len(s) <= max_chars:
        return s
    return s[: max_chars - 40] + "\n\n… (truncated) …"


def mark_dirty() -> None:
    st.session_state["_dirty"] = True


# ============================================================
# VECTOR / VOICE VAULT
# ============================================================
def _tokenize(text: str) -> List[str]:
    return [w.lower() for w in WORD_RE.findall(text or "")]


def _hash_vec(text: str, dims: int = 512) -> List[float]:
    vec = [0.0] * dims
    toks = _tokenize(text)
    for t in toks:
        h = int(hashlib.md5(t.encode("utf-8")).hexdigest(), 16)
        vec[h % dims] += 1.0
    for i, v in enumerate(vec):
        if v > 0:
            vec[i] = 1.0 + math.log(v)
    return vec


def _cosine(a: List[float], b: List[float]) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a))
    nb = math.sqrt(sum(y * y for y in b))
    if na == 0 or nb == 0:
        return 0.0
    return dot / (na * nb)


def default_voice_vault() -> Dict[str, Any]:
    ts = now_ts()
    return {
        "Voice A": {"created_ts": ts, "lanes": {ln: [] for ln in LANES}},
        "Voice B": {"created_ts": ts, "lanes": {ln: [] for ln in LANES}},
    }


def rebuild_vectors_in_voice_vault(compact_voices: Dict[str, Any]) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    for vname, v in (compact_voices or {}).items():
        created_ts = v.get("created_ts") or now_ts()
        lanes_in = v.get("lanes", {}) or {}
        lanes_out: Dict[str, Any] = {ln: [] for ln in LANES}
        for ln in LANES:
            samples = lanes_in.get(ln, []) or []
            for s in samples:
                txt = _normalize_text(s.get("text", ""))
                if not txt:
                    continue
                lanes_out[ln].append(
                    {"ts": s.get("ts") or now_ts(), "text": txt, "vec": _hash_vec(txt)}
                )
        out[vname] = {"created_ts": created_ts, "lanes": lanes_out}
    return out


def compact_voice_vault(voices: Dict[str, Any]) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    for vname, v in (voices or {}).items():
        lanes_out: Dict[str, Any] = {}
        for ln in LANES:
            samples = (v.get("lanes", {}) or {}).get(ln, []) or []
            lanes_out[ln] = [
                {"ts": s.get("ts"), "text": s.get("text", "")}
                for s in samples
                if (s.get("text") or "").strip()
            ]
        out[vname] = {"created_ts": v.get("created_ts"), "lanes": lanes_out}
    return out


def voice_names_for_selector() -> List[str]:
    names = list((st.session_state.get("voices") or {}).keys())
    names = [n for n in names if n and n.strip()]
    return ["— None —"] + sorted(set(names), key=lambda x: x.lower())


def create_my_voice_profile(voice_name: str, initial_samples: str = "") -> bool:
    """Create a new My Voice profile with optional initial training samples."""
    voice_name = (voice_name or "").strip()
    if not voice_name or voice_name == "— None —":
        return False

    if voice_name in st.session_state.voices:
        return False  # Already exists

    # Create voice profile with empty lanes
    st.session_state.voices[voice_name] = {
        "created_ts": now_ts(),
        "lanes": {ln: [] for ln in LANES},
        "my_voice": True,  # Mark as My Voice profile
        "learning_data": init_style_learning_engine(),  # Each voice has its own learning
    }

    # Add initial samples if provided
    if initial_samples:
        # Detect lane and add sample
        detected_lane = current_lane_from_draft(initial_samples)
        add_voice_sample(voice_name, detected_lane, initial_samples)

    mark_dirty()
    return True


def delete_my_voice_profile(voice_name: str) -> bool:
    """Delete a My Voice profile."""
    voice_name = (voice_name or "").strip()
    if not voice_name or voice_name == "— None —":
        return False

    if voice_name not in st.session_state.voices:
        return False

    # Don't delete default voices
    voice_data = st.session_state.voices.get(voice_name, {})
    if not voice_data.get("my_voice"):
        return False

    del st.session_state.voices[voice_name]
    mark_dirty()
    return True


def get_my_voice_profiles() -> List[str]:
    """Get list of My Voice profile names."""
    voices = st.session_state.get("voices", {})
    return [
        name
        for name, data in voices.items()
        if isinstance(data, dict) and data.get("my_voice")
    ]


def get_voice_sample_count(voice_name: str) -> int:
    """Get total number of samples for a voice."""
    voices = st.session_state.get("voices", {})
    if voice_name not in voices:
        return 0

    voice_data = voices[voice_name]
    lanes = voice_data.get("lanes", {})
    total = sum(len(samples) for samples in lanes.values())
    return total


def learn_voice_from_edit(voice_name: str, before_text: str, after_text: str) -> None:
    """Learn voice-specific patterns from an edit."""
    if not voice_name or voice_name == "— None —":
        return

    voices = st.session_state.get("voices", {})
    if voice_name not in voices:
        return

    voice_data = voices[voice_name]
    if not voice_data.get("my_voice"):
        return

    # Get or create learning data for this voice
    if "learning_data" not in voice_data:
        voice_data["learning_data"] = init_style_learning_engine()

    learning_engine = voice_data["learning_data"]

    # Add edit pair
    edit_entry = {
        "before": before_text[:500],
        "after": after_text[:500],
        "timestamp": now_ts(),
        "context": "voice_edit",
    }
    learning_engine["edit_pairs"].insert(0, edit_entry)
    learning_engine["edit_pairs"] = learning_engine["edit_pairs"][:100]

    # Update statistics
    learning_engine["style_stats"]["total_edits"] = (
        learning_engine["style_stats"].get("total_edits", 0) + 1
    )
    learning_engine["last_updated_ts"] = now_ts()

    # Learn patterns
    _extract_and_learn_patterns(after_text, learning_engine)


def get_voice_learning_stats(voice_name: str) -> Dict[str, Any]:
    """Get learning statistics for a specific voice."""
    voices = st.session_state.get("voices", {})
    if voice_name not in voices:
        return {"enabled": False, "total_edits": 0}

    voice_data = voices[voice_name]
    learning_engine = voice_data.get("learning_data")

    if not learning_engine:
        return {"enabled": False, "total_edits": 0}

    return {
        "enabled": True,
        "total_edits": learning_engine["style_stats"].get("total_edits", 0),
        "total_accepts": learning_engine["style_stats"].get("total_accepts", 0),
        "total_rejects": learning_engine["style_stats"].get("total_rejects", 0),
        "avg_sentence_length": learning_engine["style_stats"].get(
            "avg_sentence_length", 0.0
        ),
        "last_updated": learning_engine.get("last_updated_ts", "never"),
    }


def add_voice_sample(voice_name: str, lane: str, text: str) -> bool:
    voice_name = (voice_name or "").strip()
    lane = lane if lane in LANES else "Narration"
    text = _normalize_text(text)
    if not voice_name or voice_name == "— None —":
        return False
    if not text:
        return False
    if voice_name not in st.session_state.voices:
        st.session_state.voices[voice_name] = {
            "created_ts": now_ts(),
            "lanes": {ln: [] for ln in LANES},
        }
    st.session_state.voices[voice_name]["lanes"][lane].append(
        {"ts": now_ts(), "text": text, "vec": _hash_vec(text)}
    )
    mark_dirty()
    return True


# ============================================================
# STYLE BANKS
# ============================================================
def default_style_banks() -> Dict[str, Any]:
    ts = now_ts()
    return {
        s: {"created_ts": ts, "lanes": {ln: [] for ln in LANES}} for s in ENGINE_STYLES
    }


def rebuild_vectors_in_style_banks(banks: Dict[str, Any]) -> Dict[str, Any]:
    src = banks or {}
    out: Dict[str, Any] = {}
    for style in ENGINE_STYLES:
        b = (src.get(style) or {}) if isinstance(src, dict) else {}
        lanes = b.get("lanes") or {}
        new_lanes: Dict[str, List[Dict[str, Any]]] = {}
        for ln in LANES:
            samples = (lanes.get(ln) or []) if isinstance(lanes, dict) else []
            rebuilt: List[Dict[str, Any]] = []
            for it in samples:
                if isinstance(it, str):
                    t = it.strip()
                    if not t:
                        continue
                    rebuilt.append({"ts": now_ts(), "text": t, "vec": _hash_vec(t)})
                    continue
                if not isinstance(it, dict):
                    continue
                t = (it.get("text") or "").strip()
                if not t:
                    continue
                vec = it.get("vec") if isinstance(it.get("vec"), list) else None
                if not vec:
                    vec = _hash_vec(t)
                rebuilt.append({"ts": it.get("ts") or now_ts(), "text": t, "vec": vec})
            new_lanes[ln] = rebuilt
        out[style] = {"created_ts": b.get("created_ts") or now_ts(), "lanes": new_lanes}
    return out if out else default_style_banks()


def compact_style_banks(banks: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(banks, dict):
        return default_style_banks()
    out: Dict[str, Any] = {}
    for style in ENGINE_STYLES:
        b = banks.get(style) or {}
        lanes = b.get("lanes") or {}
        c_lanes: Dict[str, List[Dict[str, Any]]] = {}
        for ln in LANES:
            ss = (lanes.get(ln) or []) if isinstance(lanes, dict) else []
            cleaned: List[Dict[str, Any]] = []
            for it in ss:
                if not isinstance(it, dict):
                    continue
                t = (it.get("text") or "").strip()
                if not t:
                    continue
                cleaned.append(
                    {"ts": it.get("ts") or now_ts(), "text": _clamp_text(t, 9000)}
                )
            c_lanes[ln] = cleaned
        out[style] = {"created_ts": b.get("created_ts") or now_ts(), "lanes": c_lanes}
    return out


# ============================================================
# ADAPTIVE STYLE LEARNING ENGINE
# ============================================================
def init_style_learning_engine() -> Dict[str, Any]:
    """Initialize the adaptive style learning engine data structure."""
    return {
        "version": "1.0",
        "created_ts": now_ts(),
        "last_updated_ts": now_ts(),
        "learning_enabled": True,
        "edit_pairs": [],  # (before, after, timestamp, context)
        "accepted_suggestions": [],  # (original, suggestion, timestamp)
        "rejected_suggestions": [],  # (original, suggestion, timestamp)
        "manual_rewrites": [],  # (before, after, timestamp)
        "learned_patterns": {
            "sentence_length": {"short": 0, "medium": 0, "long": 0},
            "paragraph_structure": {"single": 0, "multi": 0},
            "word_preferences": {},  # word -> frequency
            "phrase_patterns": {},  # phrase -> frequency
            "tone_indicators": {},  # tone marker -> frequency
            "punctuation_style": {},  # pattern -> frequency
        },
        "style_stats": {
            "avg_sentence_length": 0.0,
            "avg_paragraph_length": 0.0,
            "total_edits": 0,
            "total_accepts": 0,
            "total_rejects": 0,
        },
    }


def learn_from_edit(before_text: str, after_text: str, context: str = "") -> None:
    """Learn from a user edit (before -> after)."""
    if not st.session_state.get("_style_learning_enabled", True):
        return

    if not before_text or not after_text:
        return

    engine = st.session_state.get("_style_learning_engine")
    if not engine:
        engine = init_style_learning_engine()
        st.session_state._style_learning_engine = engine

    # Add edit pair (keep last 100 edits)
    edit_entry = {
        "before": before_text[:500],
        "after": after_text[:500],
        "timestamp": now_ts(),
        "context": context,
    }
    engine["edit_pairs"].insert(0, edit_entry)
    engine["edit_pairs"] = engine["edit_pairs"][:100]

    # Update statistics
    engine["style_stats"]["total_edits"] += 1
    engine["last_updated_ts"] = now_ts()

    # Learn patterns from the edit
    _extract_and_learn_patterns(after_text, engine)


def learn_from_acceptance(original: str, suggestion: str) -> None:
    """Learn from an accepted suggestion."""
    if not st.session_state.get("_style_learning_enabled", True):
        return

    engine = st.session_state.get("_style_learning_engine")
    if not engine:
        engine = init_style_learning_engine()
        st.session_state._style_learning_engine = engine

    accept_entry = {
        "original": original[:500],
        "suggestion": suggestion[:500],
        "timestamp": now_ts(),
    }
    engine["accepted_suggestions"].insert(0, accept_entry)
    engine["accepted_suggestions"] = engine["accepted_suggestions"][:100]

    engine["style_stats"]["total_accepts"] += 1
    engine["last_updated_ts"] = now_ts()

    _extract_and_learn_patterns(suggestion, engine)


def learn_from_rejection(original: str, suggestion: str) -> None:
    """Learn from a rejected suggestion."""
    if not st.session_state.get("_style_learning_enabled", True):
        return

    engine = st.session_state.get("_style_learning_engine")
    if not engine:
        engine = init_style_learning_engine()
        st.session_state._style_learning_engine = engine

    reject_entry = {
        "original": original[:500],
        "suggestion": suggestion[:500],
        "timestamp": now_ts(),
    }
    engine["rejected_suggestions"].insert(0, reject_entry)
    engine["rejected_suggestions"] = engine["rejected_suggestions"][:100]

    engine["style_stats"]["total_rejects"] += 1
    engine["last_updated_ts"] = now_ts()


def learn_from_manual_rewrite(before_text: str, after_text: str) -> None:
    """Learn from a manual rewrite (significant user change)."""
    if not st.session_state.get("_style_learning_enabled", True):
        return

    if not before_text or not after_text:
        return

    # Only count as manual rewrite if significant change (>30% different)
    if len(before_text) > 0 and len(after_text) > 0:
        diff_ratio = abs(len(after_text) - len(before_text)) / max(
            len(before_text), len(after_text)
        )
        if diff_ratio < 0.3:
            return

    engine = st.session_state.get("_style_learning_engine")
    if not engine:
        engine = init_style_learning_engine()
        st.session_state._style_learning_engine = engine

    rewrite_entry = {
        "before": before_text[:500],
        "after": after_text[:500],
        "timestamp": now_ts(),
    }
    engine["manual_rewrites"].insert(0, rewrite_entry)
    engine["manual_rewrites"] = engine["manual_rewrites"][:100]

    engine["last_updated_ts"] = now_ts()
    _extract_and_learn_patterns(after_text, engine)


def _extract_and_learn_patterns(text: str, engine: Dict[str, Any]) -> None:
    """Extract and learn patterns from text."""
    if not text:
        return

    patterns = engine["learned_patterns"]

    # Learn sentence length preferences
    sentences = [
        s.strip()
        for s in text.replace("!", ".").replace("?", ".").split(".")
        if s.strip()
    ]
    for sent in sentences:
        word_count = len(sent.split())
        if word_count <= 10:
            patterns["sentence_length"]["short"] += 1
        elif word_count <= 20:
            patterns["sentence_length"]["medium"] += 1
        else:
            patterns["sentence_length"]["long"] += 1

    # Learn word preferences (common words)
    words = text.lower().split()
    for word in words:
        word = word.strip(".,!?;:")
        if len(word) > 3:  # Only learn significant words
            patterns["word_preferences"][word] = (
                patterns["word_preferences"].get(word, 0) + 1
            )

    # Learn punctuation style
    for char in ["!", "?", ";", ":", "--", "..."]:
        if char in text:
            patterns["punctuation_style"][char] = (
                patterns["punctuation_style"].get(char, 0) + 1
            )

    # Update averages
    if sentences:
        avg_sent_len = sum(len(s.split()) for s in sentences) / len(sentences)
        current_avg = engine["style_stats"].get("avg_sentence_length", 0.0)
        total_edits = engine["style_stats"].get("total_edits", 1)
        # Incremental average
        engine["style_stats"]["avg_sentence_length"] = (
            current_avg * (total_edits - 1) + avg_sent_len
        ) / total_edits


def get_style_learning_stats() -> Dict[str, Any]:
    """Get current style learning statistics."""
    engine = st.session_state.get("_style_learning_engine")
    if not engine:
        return {
            "enabled": True,
            "total_edits": 0,
            "total_accepts": 0,
            "total_rejects": 0,
            "avg_sentence_length": 0.0,
        }

    return {
        "enabled": st.session_state.get("_style_learning_enabled", True),
        "total_edits": engine["style_stats"].get("total_edits", 0),
        "total_accepts": engine["style_stats"].get("total_accepts", 0),
        "total_rejects": engine["style_stats"].get("total_rejects", 0),
        "avg_sentence_length": engine["style_stats"].get("avg_sentence_length", 0.0),
        "last_updated": engine.get("last_updated_ts", "never"),
        "edit_pairs_count": len(engine.get("edit_pairs", [])),
        "learned_words": len(
            engine.get("learned_patterns", {}).get("word_preferences", {})
        ),
    }


def reset_style_learning() -> None:
    """Reset the style learning engine to defaults."""
    st.session_state._style_learning_engine = init_style_learning_engine()


def export_learned_style() -> Dict[str, Any]:
    """Export the learned style for backup."""
    engine = st.session_state.get("_style_learning_engine")
    if not engine:
        return init_style_learning_engine()
    return engine.copy()


def import_learned_style(style_data: Dict[str, Any]) -> bool:
    """Import a learned style from backup."""
    if not isinstance(style_data, dict):
        return False

    if "version" not in style_data:
        return False

    st.session_state._style_learning_engine = style_data
    return True


def toggle_style_learning(enabled: bool) -> None:
    """Enable or disable style learning."""
    st.session_state._style_learning_enabled = enabled


# ============================================================
# LANE DETECTION
# ============================================================
THOUGHT_WORDS = {
    "think",
    "thought",
    "felt",
    "wondered",
    "realized",
    "remembered",
    "knew",
    "noticed",
    "decided",
    "hoped",
    "feared",
    "wanted",
    "imagined",
    "could",
    "should",
    "would",
}

ACTION_VERBS = {
    "run",
    "ran",
    "walk",
    "walked",
    "grab",
    "grabbed",
    "push",
    "pushed",
    "pull",
    "pulled",
    "slam",
    "slammed",
    "hit",
    "struck",
    "kick",
    "kicked",
    "turn",
    "turned",
    "snap",
    "snapped",
    "dive",
    "dived",
    "duck",
    "ducked",
    "rush",
    "rushed",
    "lunge",
    "lunged",
    "climb",
    "climbed",
    "drop",
    "dropped",
    "throw",
    "threw",
    "fire",
    "fired",
    "aim",
    "aimed",
    "break",
    "broke",
    "shatter",
    "shattered",
    "step",
    "stepped",
    "move",
    "moved",
    "reach",
    "reached",
}


def detect_lane(paragraph: str) -> str:
    p = (paragraph or "").strip()
    if not p:
        return "Narration"
    quote_count = p.count('"') + p.count('"') + p.count('"')
    has_dialogue_punct = p.startswith(("—", "- ", '"', '"'))
    dialogue_score = 0.0
    if quote_count >= 2:
        dialogue_score += 2.5
    if has_dialogue_punct:
        dialogue_score += 1.5
    toks = _tokenize(p)
    interior_score = 0.0
    if toks:
        first_person = sum(1 for t in toks if t in ("i", "me", "my", "mine", "myself"))
        thought_hits = sum(1 for t in toks if t in THOUGHT_WORDS)
        if first_person >= 2 and thought_hits >= 1:
            interior_score += 2.2
    action_score = 0.0
    if toks:
        verb_hits = sum(1 for t in toks if t in ACTION_VERBS)
        if verb_hits >= 2:
            action_score += 1.6
        if "!" in p:
            action_score += 0.3
    scores = {
        "Dialogue": dialogue_score,
        "Interiority": interior_score,
        "Action": action_score,
        "Narration": 0.25,
    }
    lane = max(scores.items(), key=lambda kv: kv[1])[0]
    return "Narration" if scores[lane] < 0.9 else lane


def current_lane_from_draft(text: str) -> str:
    paras = _split_paragraphs(text)
    if not paras:
        return "Narration"
    return detect_lane(paras[-1])


# ============================================================
# INTENSITY
# ============================================================
def intensity_profile(x: float) -> str:
    if x <= 0.25:
        return "LOW: conservative, literal, minimal invention, prioritize continuity and clarity."
    if x <= 0.60:
        return "MED: balanced creativity, stronger phrasing, controlled invention within canon."
    if x <= 0.85:
        return "HIGH: bolder choices, richer specificity; still obey canon and lane."
    return "MAX: aggressive originality and voice; still obey canon, no derailments."


def temperature_from_intensity(x: float) -> float:
    x = max(0.0, min(1.0, float(x)))
    return 0.15 + (x * 0.95)


# ============================================================
# STORY BIBLE WORKSPACE
# ============================================================
def default_story_bible_workspace() -> Dict[str, Any]:
    ts = now_ts()
    return {
        "workspace_story_bible_id": hashlib.md5(
            f"wsb|{ts}".encode("utf-8")
        ).hexdigest()[:12],
        "workspace_story_bible_created_ts": ts,
        "title": "",
        "draft": "",
        "story_bible": {
            "synopsis": "",
            "genre_style_notes": "",
            "world": "",
            "characters": "",
            "outline": "",
        },
        "voice_sample": "",
        "ai_intensity": 0.75,
        "voices": default_voice_vault(),
        "style_banks": default_style_banks(),
    }


def in_workspace_mode() -> bool:
    return (st.session_state.active_bay == "NEW") and (
        st.session_state.project_id is None
    )


def save_workspace_from_session() -> None:
    w = st.session_state.sb_workspace or default_story_bible_workspace()
    w["title"] = st.session_state.get("workspace_title", w.get("title", ""))
    w["draft"] = st.session_state.main_text
    w["story_bible"] = {
        "synopsis": st.session_state.synopsis,
        "genre_style_notes": st.session_state.genre_style_notes,
        "world": st.session_state.world,
        "characters": st.session_state.characters,
        "outline": st.session_state.outline,
    }
    w["voice_sample"] = st.session_state.voice_sample
    w["ai_intensity"] = float(st.session_state.ai_intensity)
    w["voices"] = compact_voice_vault(st.session_state.voices)
    w["style_banks"] = compact_style_banks(
        st.session_state.get("style_banks")
        or rebuild_vectors_in_style_banks(default_style_banks())
    )
    st.session_state.sb_workspace = w


def load_workspace_into_session() -> None:
    w = st.session_state.sb_workspace or default_story_bible_workspace()
    sb = w.get("story_bible", {}) or {}
    st.session_state.project_id = None
    st.session_state.project_title = "—"
    st.session_state.main_text = w.get("draft", "") or ""
    # Initialize undo history with loaded text
    st.session_state._undo_history = [st.session_state.main_text]
    st.session_state._redo_history = []
    st.session_state.synopsis = sb.get("synopsis", "") or ""
    st.session_state.genre_style_notes = sb.get("genre_style_notes", "") or ""
    st.session_state.world = sb.get("world", "") or ""
    st.session_state.characters = sb.get("characters", "") or ""
    st.session_state.outline = sb.get("outline", "") or ""
    st.session_state.voice_sample = w.get("voice_sample", "") or ""
    st.session_state.ai_intensity = float(w.get("ai_intensity", 0.75))
    st.session_state.voices = rebuild_vectors_in_voice_vault(
        w.get("voices", default_voice_vault())
    )
    st.session_state.style_banks = rebuild_vectors_in_style_banks(
        w.get("style_banks", default_style_banks())
    )
    st.session_state.workspace_title = w.get("title", "") or ""


def reset_workspace_story_bible(keep_templates: bool = True) -> None:
    old = st.session_state.sb_workspace or default_story_bible_workspace()
    neww = default_story_bible_workspace()
    if keep_templates:
        neww["voice_sample"] = old.get("voice_sample", "")
        neww["ai_intensity"] = float(old.get("ai_intensity", 0.75))
        neww["voices"] = old.get("voices", default_voice_vault())
        neww["style_banks"] = old.get("style_banks", default_style_banks())
    st.session_state.sb_workspace = neww
    if in_workspace_mode():
        load_workspace_into_session()
    mark_dirty()


# ============================================================
# CLOUD SYNC SUPPORT (S3-Compatible)
# ============================================================
def get_s3_credentials() -> Optional[Dict[str, str]]:
    """Get S3 credentials from secrets or environment."""
    try:
        # Try Streamlit secrets first
        if hasattr(st, "secrets"):
            aws_key = st.secrets.get("AWS_ACCESS_KEY_ID", "")
            aws_secret = st.secrets.get("AWS_SECRET_ACCESS_KEY", "")
            aws_bucket = st.secrets.get("AWS_S3_BUCKET", "")
            aws_region = st.secrets.get("AWS_REGION", "us-east-1")
            aws_endpoint = st.secrets.get("AWS_ENDPOINT_URL", "")  # For MinIO/custom
        else:
            aws_key = ""
            aws_secret = ""
            aws_bucket = ""
            aws_region = ""
            aws_endpoint = ""

        # Fall back to environment variables
        aws_key = aws_key or os.getenv("AWS_ACCESS_KEY_ID", "")
        aws_secret = aws_secret or os.getenv("AWS_SECRET_ACCESS_KEY", "")
        aws_bucket = aws_bucket or os.getenv("AWS_S3_BUCKET", "")
        aws_region = aws_region or os.getenv("AWS_REGION", "us-east-1")
        aws_endpoint = aws_endpoint or os.getenv("AWS_ENDPOINT_URL", "")

        if not (aws_key and aws_secret and aws_bucket):
            return None

        return {
            "aws_access_key_id": aws_key,
            "aws_secret_access_key": aws_secret,
            "bucket": aws_bucket,
            "region": aws_region,
            "endpoint_url": aws_endpoint if aws_endpoint else None,
        }
    except Exception:
        return None


def has_cloud_sync() -> bool:
    """Check if cloud sync is configured."""
    return get_s3_credentials() is not None


def list_cloud_saves() -> List[Dict[str, Any]]:
    """List available saves in cloud storage."""
    try:
        import boto3
        from botocore.exceptions import ClientError

        creds = get_s3_credentials()
        if not creds:
            return []

        # Create S3 client
        s3_config = {
            "aws_access_key_id": creds["aws_access_key_id"],
            "aws_secret_access_key": creds["aws_secret_access_key"],
            "region_name": creds["region"],
        }
        if creds["endpoint_url"]:
            s3_config["endpoint_url"] = creds["endpoint_url"]

        s3 = boto3.client("s3", **s3_config)

        # List objects with prefix
        response = s3.list_objects_v2(Bucket=creds["bucket"], Prefix="olivetti/")

        saves = []
        for obj in response.get("Contents", []):
            key = obj["Key"]
            if key.endswith(".json"):
                save_name = key.replace("olivetti/", "").replace(".json", "")
                saves.append(
                    {
                        "name": save_name,
                        "key": key,
                        "size": obj["Size"],
                        "modified": obj["LastModified"].strftime("%Y-%m-%d %H:%M:%S"),
                    }
                )

        # Sort by modified date (newest first)
        saves.sort(key=lambda x: x["modified"], reverse=True)
        return saves
    except ImportError:
        st.error("boto3 not installed. Run: pip install boto3")
        return []
    except Exception as e:
        st.error(f"Failed to list cloud saves: {e}")
        return []


def upload_to_cloud(save_name: str) -> bool:
    """Upload current state to cloud storage."""
    try:
        import boto3
        from botocore.exceptions import ClientError

        creds = get_s3_credentials()
        if not creds:
            st.error("Cloud sync not configured")
            return False

        # Create snapshot
        snapshot = {
            "saved_ts": now_ts(),
            "save_name": save_name,
            "session": {
                k: (
                    compact_voice_vault(v)
                    if k == "voices"
                    else (
                        compact_style_banks(v)
                        if k == "style_banks"
                        else st.session_state.get(k)
                    )
                )
                for k, v in st.session_state.items()
                if k not in ("_rerun_count",)
            },
        }

        # Serialize to JSON
        json_data = json.dumps(snapshot, ensure_ascii=False, indent=2)

        # Upload to S3
        s3_config = {
            "aws_access_key_id": creds["aws_access_key_id"],
            "aws_secret_access_key": creds["aws_secret_access_key"],
            "region_name": creds["region"],
        }
        if creds["endpoint_url"]:
            s3_config["endpoint_url"] = creds["endpoint_url"]

        s3 = boto3.client("s3", **s3_config)

        key = f"olivetti/{save_name}.json"
        s3.put_object(
            Bucket=creds["bucket"],
            Key=key,
            Body=json_data.encode("utf-8"),
            ContentType="application/json",
        )

        return True
    except ImportError:
        st.error("boto3 not installed. Run: pip install boto3")
        return False
    except Exception as e:
        st.error(f"Failed to upload to cloud: {e}")
        return False


def download_from_cloud(save_key: str) -> bool:
    """Download state from cloud storage."""
    try:
        import boto3
        from botocore.exceptions import ClientError

        creds = get_s3_credentials()
        if not creds:
            st.error("Cloud sync not configured")
            return False

        # Download from S3
        s3_config = {
            "aws_access_key_id": creds["aws_access_key_id"],
            "aws_secret_access_key": creds["aws_secret_access_key"],
            "region_name": creds["region"],
        }
        if creds["endpoint_url"]:
            s3_config["endpoint_url"] = creds["endpoint_url"]

        s3 = boto3.client("s3", **s3_config)

        response = s3.get_object(
            Bucket=creds["bucket"],
            Key=save_key,
        )

        # Parse JSON
        json_data = response["Body"].read().decode("utf-8")
        snapshot = json.loads(json_data)

        # Load session state
        sess = snapshot.get("session", {}) or {}
        for k in (
            "voices",
            "style_banks",
            "sb_workspace",
            "main_text",
            "synopsis",
            "genre_style_notes",
            "world",
            "characters",
            "outline",
            "ai_intensity",
            "workspace_title",
            "story_bible_lock",
        ):
            if k not in sess:
                continue
            if k == "voices":
                st.session_state.voices = rebuild_vectors_in_voice_vault(
                    sess.get("voices", default_voice_vault())
                )
            elif k == "style_banks":
                st.session_state.style_banks = rebuild_vectors_in_style_banks(
                    sess.get("style_banks", default_style_banks())
                )
            else:
                st.session_state[k] = sess.get(k)

        # Reset undo history after cloud load
        st.session_state._undo_history = [st.session_state.get("main_text", "")]
        st.session_state._redo_history = []
        st.session_state["_dirty"] = False

        return True
    except ImportError:
        st.error("boto3 not installed. Run: pip install boto3")
        return False
    except Exception as e:
        st.error(f"Failed to download from cloud: {e}")
        return False


def delete_from_cloud(save_key: str) -> bool:
    """Delete a save from cloud storage."""
    try:
        import boto3
        from botocore.exceptions import ClientError

        creds = get_s3_credentials()
        if not creds:
            st.error("Cloud sync not configured")
            return False

        s3_config = {
            "aws_access_key_id": creds["aws_access_key_id"],
            "aws_secret_access_key": creds["aws_secret_access_key"],
            "region_name": creds["region"],
        }
        if creds["endpoint_url"]:
            s3_config["endpoint_url"] = creds["endpoint_url"]

        s3 = boto3.client("s3", **s3_config)

        s3.delete_object(
            Bucket=creds["bucket"],
            Key=save_key,
        )

        return True
    except ImportError:
        st.error("boto3 not installed. Run: pip install boto3")
        return False
    except Exception as e:
        st.error(f"Failed to delete from cloud: {e}")
        return False


# ============================================================
# IMPORT / EXPORT SUPPORT
# ============================================================
def import_text_from_txt(uploaded_file) -> Optional[str]:
    """Import text from TXT file."""
    try:
        content = uploaded_file.read().decode("utf-8")
        return _normalize_text(content)
    except Exception as e:
        st.error(f"Failed to import TXT: {e}")
        return None


def import_text_from_md(uploaded_file) -> Optional[str]:
    """Import text from Markdown file."""
    try:
        content = uploaded_file.read().decode("utf-8")
        return _normalize_text(content)
    except Exception as e:
        st.error(f"Failed to import MD: {e}")
        return None


def import_text_from_docx(uploaded_file) -> Optional[str]:
    """Import text from DOCX file using python-docx."""
    try:
        # Import dynamically to avoid hard dependency
        import docx

        doc = docx.Document(uploaded_file)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)
        return _normalize_text(content)
    except ImportError:
        st.error("python-docx not installed. Run: pip install python-docx")
        return None
    except Exception as e:
        st.error(f"Failed to import DOCX: {e}")
        return None


def import_text_from_pdf(uploaded_file) -> Optional[str]:
    """Import text from PDF file using pypdf."""
    try:
        # Import dynamically to avoid hard dependency
        import pypdf

        reader = pypdf.PdfReader(uploaded_file)
        paragraphs = []
        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                paragraphs.append(text.strip())
        content = "\n\n".join(paragraphs)
        return _normalize_text(content)
    except ImportError:
        st.error("pypdf not installed. Run: pip install pypdf")
        return None
    except Exception as e:
        st.error(f"Failed to import PDF: {e}")
        return None


def export_to_txt(text: str, filename: str) -> bytes:
    """Export text to TXT format."""
    return text.encode("utf-8")


def export_to_md(text: str, filename: str) -> bytes:
    """Export text to Markdown format."""
    # Add a title if we have one
    title = st.session_state.get("workspace_title", "")
    if title:
        md_content = f"# {title}\n\n{text}"
    else:
        md_content = text
    return md_content.encode("utf-8")


def export_to_docx(text: str, filename: str) -> Optional[bytes]:
    """Export text to DOCX format using python-docx."""
    try:
        import docx
        from io import BytesIO

        doc = docx.Document()

        # Add title if available
        title = st.session_state.get("workspace_title", "")
        if title:
            doc.add_heading(title, level=0)

        # Add paragraphs
        paragraphs = _split_paragraphs(text)
        for para in paragraphs:
            doc.add_paragraph(para)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    except ImportError:
        st.error("python-docx not installed. Run: pip install python-docx")
        return None
    except Exception as e:
        st.error(f"Failed to export DOCX: {e}")
        return None


def export_to_pdf(text: str, filename: str) -> Optional[bytes]:
    """Export text to PDF format using reportlab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from io import BytesIO

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        styles = getSampleStyleSheet()
        story = []

        # Add title if available
        title = st.session_state.get("workspace_title", "")
        if title:
            title_style = styles["Title"]
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.2 * inch))

        # Add paragraphs
        paragraphs = _split_paragraphs(text)
        body_style = styles["BodyText"]
        for para in paragraphs:
            # Escape XML special characters
            para_escaped = (
                para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            )
            story.append(Paragraph(para_escaped, body_style))
            story.append(Spacer(1, 0.15 * inch))

        doc.build(story)
        buffer.seek(0)
        return buffer.read()
    except ImportError:
        st.error("reportlab not installed. Run: pip install reportlab")
        return None
    except Exception as e:
        st.error(f"Failed to export PDF: {e}")
        return None


# ============================================================
# BAY TRANSFER & EXPORT
# ============================================================
def can_transfer_bay() -> bool:
    """Check if current bay can be transferred to next bay."""
    current_bay = st.session_state.get("active_bay", "NEW")
    # Can transfer from ROUGH to EDIT, or EDIT to FINAL
    return current_bay in ["ROUGH", "EDIT"]


def get_next_bay() -> Optional[str]:
    """Get the next bay in the workflow."""
    current_bay = st.session_state.get("active_bay", "NEW")
    if current_bay == "ROUGH":
        return "EDIT"
    elif current_bay == "EDIT":
        return "FINAL"
    return None


def transfer_to_next_bay() -> bool:
    """Transfer current draft to the next bay in the workflow."""
    if not can_transfer_bay():
        return False

    current_bay = st.session_state.get("active_bay", "NEW")
    next_bay = get_next_bay()

    if not next_bay:
        return False

    # Check if next bay already has content
    next_project = st.session_state.active_project_by_bay.get(next_bay)
    if next_project:
        # Next bay has content - need confirmation
        st.session_state._transfer_needs_confirmation = True
        st.session_state._transfer_target_bay = next_bay
        return False

    # Perform transfer
    _execute_bay_transfer(next_bay)
    return True


def _execute_bay_transfer(target_bay: str) -> None:
    """Execute the bay transfer (internal helper)."""
    current_bay = st.session_state.get("active_bay", "NEW")

    # Save current state to workspace
    save_workspace_from_session()

    # Create a copy of the current workspace for the target bay
    current_workspace = st.session_state.sb_workspace.copy()

    # Store in the target bay
    st.session_state.active_project_by_bay[target_bay] = current_workspace

    # Switch to the target bay
    st.session_state.active_bay = target_bay

    # Load the workspace (which is the same content we just transferred)
    load_workspace_into_session()

    mark_dirty()


def get_export_filename() -> str:
    """Generate export filename based on workspace title and bay."""
    title = st.session_state.get("workspace_title", "")
    bay = st.session_state.get("active_bay", "NEW")

    if title:
        base = re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "_")
        return f"{base}_{bay.lower()}"
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"olivetti_{bay.lower()}_{timestamp}"


def count_words(text: str) -> int:
    """Count words in text."""
    if not text or not text.strip():
        return 0
    return len(text.split())


def get_voice_settings() -> dict:
    """Get current voice settings from session state."""
    return {
        "vb_style_on": st.session_state.get("vb_style_on", True),
        "writing_style": st.session_state.get("writing_style", "Neutral"),
        "ai_intensity": st.session_state.get("ai_intensity", 0.75),
        "trained_voice": st.session_state.get("trained_voice", "None"),
        "voice_lane": st.session_state.get("voice_lane", "Narration"),
        "voice_sample": st.session_state.get("voice_sample", ""),
    }


def set_voice_settings(settings: dict) -> None:
    """Set voice settings in session state."""
    st.session_state.vb_style_on = settings.get("vb_style_on", True)
    st.session_state.writing_style = settings.get("writing_style", "Neutral")
    st.session_state.ai_intensity = settings.get("ai_intensity", 0.75)
    st.session_state.trained_voice = settings.get("trained_voice", "None")
    st.session_state.voice_lane = settings.get("voice_lane", "Narration")
    st.session_state.voice_sample = settings.get("voice_sample", "")


def save_voice_settings_for_bay(bay: str) -> None:
    """Save current voice settings for a specific bay."""
    if "_voice_settings_by_bay" not in st.session_state:
        st.session_state._voice_settings_by_bay = {}

    st.session_state._voice_settings_by_bay[bay] = get_voice_settings()


def load_voice_settings_for_bay(bay: str) -> None:
    """Load voice settings for a specific bay."""
    voice_settings_by_bay = st.session_state.get("_voice_settings_by_bay", {})

    if bay in voice_settings_by_bay:
        # Load saved settings for this bay
        set_voice_settings(voice_settings_by_bay[bay])
    else:
        # Use default settings for this bay
        default_settings = {
            "vb_style_on": True,
            "writing_style": "Neutral",
            "ai_intensity": 0.75,
            "trained_voice": "None",
            "voice_lane": "Narration",
            "voice_sample": "",
        }
        set_voice_settings(default_settings)


def on_bay_change() -> None:
    """Callback when active bay changes - save old bay settings and load new bay settings."""
    # Get the previous bay from a temporary storage
    previous_bay = st.session_state.get("_previous_bay")
    current_bay = st.session_state.get("active_bay", "NEW")

    # Save voice settings for the previous bay
    if previous_bay and previous_bay != current_bay:
        save_voice_settings_for_bay(previous_bay)

    # Load voice settings for the new bay
    load_voice_settings_for_bay(current_bay)

    # Update the previous bay tracker
    st.session_state._previous_bay = current_bay

    mark_dirty()


def delete_current_draft() -> bool:
    """Delete the current draft from the active bay (moves to trash bin)."""
    current_bay = st.session_state.get("active_bay", "NEW")

    # Check if there's content to delete
    current_text = st.session_state.get("main_text", "").strip()
    if not current_text:
        return False

    # Save current workspace state to trash before deleting
    save_workspace_from_session()
    deleted_workspace = st.session_state.sb_workspace.copy()

    # Initialize trash bin if needed
    if "_trash_bin" not in st.session_state or st.session_state._trash_bin is None:
        st.session_state._trash_bin = []

    # Add to trash bin (newest first)
    trash_item = {
        "bay": current_bay,
        "workspace": deleted_workspace,
        "timestamp": datetime.now().isoformat(),
        "id": f"{current_bay}_{datetime.now().timestamp()}",
    }
    st.session_state._trash_bin.insert(0, trash_item)

    # Keep only the 10 most recent items
    MAX_TRASH_ITEMS = 10
    if len(st.session_state._trash_bin) > MAX_TRASH_ITEMS:
        st.session_state._trash_bin = st.session_state._trash_bin[:MAX_TRASH_ITEMS]

    # Clear the current session state
    st.session_state.main_text = ""
    st.session_state.workspace_title = ""
    st.session_state._undo_history = [""]
    st.session_state._redo_history = []

    # Remove from active_project_by_bay if it exists
    if current_bay in st.session_state.active_project_by_bay:
        del st.session_state.active_project_by_bay[current_bay]

    mark_dirty()
    return True


def get_trash_bin_items() -> list:
    """Get all items in the trash bin."""
    trash_bin = st.session_state.get("_trash_bin", [])
    return trash_bin if trash_bin else []


def restore_from_trash(trash_id: str) -> bool:
    """Restore a specific item from the trash bin."""
    trash_bin = get_trash_bin_items()

    # Find the item
    item_to_restore = None
    item_index = None
    for i, item in enumerate(trash_bin):
        if item.get("id") == trash_id:
            item_to_restore = item
            item_index = i
            break

    if not item_to_restore:
        return False

    deleted_bay = item_to_restore["bay"]
    deleted_workspace = item_to_restore["workspace"]

    # Restore the workspace to its original bay
    st.session_state.active_project_by_bay[deleted_bay] = deleted_workspace

    # Switch to the restored bay
    st.session_state.active_bay = deleted_bay

    # Load the restored workspace into session
    load_workspace_into_session()

    # Remove from trash bin
    st.session_state._trash_bin.pop(item_index)

    mark_dirty()
    return True


def permanently_delete_from_trash(trash_id: str) -> bool:
    """Permanently delete a specific item from the trash bin."""
    trash_bin = get_trash_bin_items()

    # Find and remove the item
    for i, item in enumerate(trash_bin):
        if item.get("id") == trash_id:
            st.session_state._trash_bin.pop(i)
            return True

    return False


def clear_trash_bin() -> None:
    """Clear all items from the trash bin."""
    st.session_state._trash_bin = []


# ============================================================
# BAY TRANSFER & EXPORT
# ============================================================
def can_transfer_bay() -> bool:
    """Check if current bay can be transferred to next bay."""
    current_bay = st.session_state.get("active_bay", "NEW")
    # Can transfer from ROUGH to EDIT, or EDIT to FINAL
    return current_bay in ["ROUGH", "EDIT"]


def get_next_bay() -> Optional[str]:
    """Get the next bay in the workflow."""
    current_bay = st.session_state.get("active_bay", "NEW")
    if current_bay == "ROUGH":
        return "EDIT"
    elif current_bay == "EDIT":
        return "FINAL"
    return None


def transfer_to_next_bay() -> bool:
    """Transfer current draft to the next bay in the workflow."""
    if not can_transfer_bay():
        return False

    current_bay = st.session_state.get("active_bay", "NEW")
    next_bay = get_next_bay()

    if not next_bay:
        return False

    # Check if next bay already has content
    next_project = st.session_state.active_project_by_bay.get(next_bay)
    if next_project:
        # Next bay has content - need confirmation
        st.session_state._transfer_needs_confirmation = True
        st.session_state._transfer_target_bay = next_bay
        return False

    # Perform transfer
    _execute_bay_transfer(next_bay)
    return True


def _execute_bay_transfer(target_bay: str) -> None:
    """Execute the bay transfer (internal helper)."""
    current_bay = st.session_state.get("active_bay", "NEW")

    # Save current state to workspace
    save_workspace_from_session()

    # Create a copy of the current workspace for the target bay
    current_workspace = st.session_state.sb_workspace.copy()

    # Store in the target bay
    st.session_state.active_project_by_bay[target_bay] = current_workspace

    # Switch to the target bay
    st.session_state.active_bay = target_bay

    # Load the workspace (which is the same content we just transferred)
    load_workspace_into_session()

    mark_dirty()


def get_export_filename() -> str:
    """Generate export filename based on workspace title and bay."""
    title = st.session_state.get("workspace_title", "")
    bay = st.session_state.get("active_bay", "NEW")

    if title:
        base = re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "_")
        return f"{base}_{bay.lower()}"
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"olivetti_{bay.lower()}_{timestamp}"


# ============================================================
# UNDO / REDO SUPPORT
# ============================================================
def push_undo_history(text: str) -> None:
    """Push current text to undo history stack."""
    history = st.session_state.get("_undo_history", [])
    # Avoid duplicate consecutive entries
    if not history or history[-1] != text:
        history.append(text)
        # Keep history bounded to 50 states
        if len(history) > 50:
            history.pop(0)
        st.session_state._undo_history = history
        # Clear redo stack when new change is made
        st.session_state._redo_history = []


def undo_text() -> None:
    """Undo last text change."""
    undo_history = st.session_state.get("_undo_history", [])
    if len(undo_history) > 1:
        # Pop current state and save to redo
        current = undo_history.pop()
        redo_history = st.session_state.get("_redo_history", [])
        redo_history.append(current)
        st.session_state._redo_history = redo_history
        # Restore previous state
        st.session_state.main_text = undo_history[-1]
        st.session_state._undo_history = undo_history
        mark_dirty()


def redo_text() -> None:
    """Redo previously undone text change."""
    redo_history = st.session_state.get("_redo_history", [])
    if redo_history:
        # Pop from redo and restore
        text = redo_history.pop()
        undo_history = st.session_state.get("_undo_history", [])
        undo_history.append(text)
        st.session_state._undo_history = undo_history
        st.session_state._redo_history = redo_history
        st.session_state.main_text = text
        mark_dirty()


def can_undo() -> bool:
    """Check if undo is available."""
    return len(st.session_state.get("_undo_history", [])) > 1


def can_redo() -> bool:
    """Check if redo is available."""
    return len(st.session_state.get("_redo_history", [])) > 0


def on_main_text_change() -> None:
    """Callback for main text changes to track history and learn from edits."""
    current_text = st.session_state.main_text

    # Get previous text from history for learning
    undo_history = st.session_state.get("_undo_history", [])
    if undo_history:
        previous_text = undo_history[-1] if undo_history else ""
        # Learn from the edit if it's significant
        if previous_text and current_text and previous_text != current_text:
            # Learn globally
            learn_from_edit(previous_text, current_text, context="main_text")

            # Learn for the selected My Voice profile
            trained_voice = st.session_state.get("trained_voice", "— None —")
            if trained_voice and trained_voice != "— None —":
                learn_voice_from_edit(trained_voice, previous_text, current_text)

    push_undo_history(current_text)
    mark_dirty()


# ============================================================
# SESSION INIT
# ============================================================
def init_state() -> None:
    defaults: Dict[str, Any] = {
        "active_bay": "NEW",
        "projects": {},
        "active_project_by_bay": {b: None for b in BAYS},
        "sb_workspace": default_story_bible_workspace(),
        "workspace_title": "",
        "project_id": None,
        "project_title": "—",
        "autosave_time": None,
        "_autosave_unix": 0.0,
        "_dirty": False,
        "last_action": "—",
        "main_text": "",
        "synopsis": "",
        "genre_style_notes": "",
        "world": "",
        "characters": "",
        "outline": "",
        "tool_output": "",
        "vb_style_on": True,
        "writing_style": "Neutral",
        "ai_intensity": 0.75,
        "trained_voice": "None",
        "voice_lane": "Narration",
        "voice_sample": "",
        "voices": rebuild_vectors_in_voice_vault(default_voice_vault()),
        "style_banks": rebuild_vectors_in_style_banks(default_style_banks()),
        "story_bible_lock": True,
        "_undo_history": [""],
        "_redo_history": [],
        "_autosave_checked": False,
        "_show_recovery_dialog": False,
        "_confirm_import": False,
        "_show_cloud_upload": False,
        "_show_cloud_download": False,
        "_cloud_save_name": "",
        "_selected_cloud_save": None,
        "_transfer_needs_confirmation": False,
        "_transfer_target_bay": None,
        "_show_export_dialog": False,
        "_confirm_delete": False,
        "_trash_bin": [],
        "_show_trash_bin": False,
        "_confirm_permanent_delete": None,
        "_voice_settings_by_bay": {},
        "_previous_bay": "NEW",
        "_style_learning_engine": init_style_learning_engine(),
        "_style_learning_enabled": True,
        "_show_style_patterns": False,
        "_confirm_reset_learning": False,
        "_confirm_delete_voice": None,
        "_system_contract": init_system_contract(),
        "_show_contract_panel": False,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


# ============================================================
# AUTOSAVE / PERSISTENCE
# ============================================================
def ensure_autosave_dir() -> None:
    if not os.path.exists(AUTOSAVE_DIR):
        os.makedirs(AUTOSAVE_DIR, exist_ok=True)


def autosave_state() -> None:
    ensure_autosave_dir()
    snapshot = {
        "saved_ts": now_ts(),
        "session": {
            k: (
                compact_voice_vault(v)
                if k == "voices"
                else (
                    compact_style_banks(v)
                    if k == "style_banks"
                    else st.session_state.get(k)
                )
            )
            for k, v in st.session_state.items()
            if k not in ("_rerun_count",)
        },
    }
    with open(AUTOSAVE_PATH, "w", encoding="utf-8") as f:
        json.dump(snapshot, f, ensure_ascii=False, indent=2)
    st.session_state.autosave_time = snapshot["saved_ts"]
    st.session_state["_autosave_unix"] = now_unix()
    st.session_state["_dirty"] = False


def maybe_autosave_throttled() -> None:
    if not st.session_state.get("_dirty"):
        return
    last = float(st.session_state.get("_autosave_unix", 0.0) or 0.0)
    if now_unix() - last < AUTOSAVE_MIN_INTERVAL_S:
        return
    autosave_state()


def get_autosave_info() -> Optional[Dict[str, Any]]:
    """Get autosave details without loading it."""
    if not os.path.exists(AUTOSAVE_PATH):
        return None
    try:
        with open(AUTOSAVE_PATH, "r", encoding="utf-8") as f:
            snap = json.load(f)
        sess = snap.get("session", {}) or {}
        main_text = sess.get("main_text", "")
        word_count = len(WORD_RE.findall(main_text)) if main_text else 0
        return {
            "saved_ts": snap.get("saved_ts", "Unknown"),
            "word_count": word_count,
            "preview": main_text[:100] + "..." if len(main_text) > 100 else main_text,
            "has_content": bool(main_text and main_text.strip()),
        }
    except Exception:
        return None


def load_autosave() -> bool:
    if not os.path.exists(AUTOSAVE_PATH):
        return False
    try:
        with open(AUTOSAVE_PATH, "r", encoding="utf-8") as f:
            snap = json.load(f)
        sess = snap.get("session", {}) or {}
        for k in (
            "voices",
            "style_banks",
            "sb_workspace",
            "main_text",
            "synopsis",
            "genre_style_notes",
            "world",
            "characters",
            "outline",
            "ai_intensity",
            "workspace_title",
            "story_bible_lock",
        ):
            if k not in sess:
                continue
            if k == "voices":
                st.session_state.voices = rebuild_vectors_in_voice_vault(
                    sess.get("voices", default_voice_vault())
                )
            elif k == "style_banks":
                st.session_state.style_banks = rebuild_vectors_in_style_banks(
                    sess.get("style_banks", default_style_banks())
                )
            else:
                st.session_state[k] = sess.get(k)
        st.session_state.autosave_time = snap.get("saved_ts", now_ts())
        st.session_state["_autosave_unix"] = now_unix()
        st.session_state["_dirty"] = False
        # Reset undo history after load
        st.session_state._undo_history = [st.session_state.get("main_text", "")]
        st.session_state._redo_history = []
        return True
    except Exception:
        return False


# ============================================================
# OPENAI CALLS (optional, safe fallback)
# ============================================================
def _openai_system_prompt(action: str, lane: str) -> str:
    sb = _normalize_text(st.session_state.synopsis)
    notes = _normalize_text(st.session_state.genre_style_notes)
    world = _normalize_text(st.session_state.world)
    chars = _normalize_text(st.session_state.characters)
    outline = _normalize_text(st.session_state.outline)
    style_on = bool(st.session_state.vb_style_on)
    style = st.session_state.writing_style if style_on else "Neutral"
    intensity = float(st.session_state.ai_intensity)
    profile = intensity_profile(intensity)
    prompt_text = f"""You are Olivetti Desk, a writing assistant.

Task: {action}
Lane focus: {lane}
Intensity: {profile}
Writing style: {style}

Rules:
- Preserve continuity and canon.
- Respect the lane (Dialogue/Narration/Interiority/Action).
- Output only the rewritten/continued prose, no commentary.

Story bible:
Synopsis: {sb}
Genre/style notes: {notes}
World: {world}
Characters: {chars}
Outline: {outline}
"""
    return _normalize_text(prompt_text.strip())


def call_openai_text(user_prompt: str, system_prompt: str, temperature: float) -> str:
    if not has_openai_key():
        return "OpenAI key not configured. Set OPENAI_API_KEY to enable model features."
    key = require_openai_key()
    # Try the modern "OpenAI-compatible" Responses API over HTTPS (no extra deps besides httpx).
    try:
        import httpx  # type: ignore

        headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}
        payload = {
            "model": OPENAI_MODEL,
            "temperature": float(temperature),
            "input": [
                {
                    "role": "system",
                    "content": [{"type": "text", "text": system_prompt}],
                },
                {"role": "user", "content": [{"type": "text", "text": user_prompt}]},
            ],
        }
        with httpx.Client(timeout=60.0) as client:
            r = client.post(
                "https://api.openai.com/v1/responses", headers=headers, json=payload
            )
            r.raise_for_status()
            data = r.json()
            # Extract text from response output
            out_texts: List[str] = []
            for item in data.get("output") or []:
                for c in item.get("content") or []:
                    if c.get("type") == "output_text":
                        out_texts.append(c.get("text", ""))
            text = _normalize_text("\n".join(out_texts))
            return text or "(No text returned.)"
    except Exception as e:
        return f"Model call failed. (Details: {e})"


def do_action(action: str) -> None:
    draft = st.session_state.main_text or ""
    lane = current_lane_from_draft(draft)
    system_prompt = _openai_system_prompt(action=action, lane=lane)
    temperature = temperature_from_intensity(float(st.session_state.ai_intensity))
    if action == "WRITE":
        user_prompt = f"Continue this draft naturally:\n\n{draft}".strip()
    elif action == "REWRITE":
        user_prompt = f"Rewrite this to improve clarity, rhythm, and voice while preserving meaning:\n\n{draft}".strip()
    elif action == "EXPAND":
        user_prompt = f"Expand this with richer specificity and sensory detail while preserving intent:\n\n{draft}".strip()
    else:
        user_prompt = draft
    st.session_state.last_action = action
    st.session_state.tool_output = call_openai_text(
        user_prompt=user_prompt, system_prompt=system_prompt, temperature=temperature
    )
    mark_dirty()


# ============================================================
# UI
# ============================================================
def main_ui() -> None:
    st.sidebar.title("Olivetti Desk")
    st.sidebar.markdown("**Workspace**")
    st.sidebar.selectbox(
        "Active Bay",
        BAYS,
        index=BAYS.index(st.session_state.active_bay),
        key="active_bay",
        on_change=on_bay_change,
    )

    # Autosave controls with recovery dialog
    autosave_info = get_autosave_info()

    c1, c2 = st.sidebar.columns(2)
    with c1:
        if st.button("Recover autosave", disabled=autosave_info is None):
            st.session_state._show_recovery_dialog = True
    with c2:
        if st.button("Save now"):
            autosave_state()
            st.sidebar.success("Saved")

    # Show autosave status
    if autosave_info:
        st.sidebar.caption(
            f"💾 Last autosave: {autosave_info['saved_ts']} "
            f"({autosave_info['word_count']} words)"
        )
    else:
        st.sidebar.caption("💾 No autosave available")

    # Recovery dialog
    if st.session_state.get("_show_recovery_dialog") and autosave_info:
        with st.sidebar.expander("⚠️ Autosave Recovery", expanded=True):
            st.warning("Loading autosave will replace your current work!")
            st.markdown("**Autosave details:**")
            st.text(f"Saved: {autosave_info['saved_ts']}")
            st.text(f"Words: {autosave_info['word_count']}")
            if autosave_info["preview"]:
                st.text("Preview:")
                st.code(autosave_info["preview"], language=None)

            col1, col2 = st.columns(2)
            with col1:
                if st.button("✓ Load it", key="confirm_recovery"):
                    ok = load_autosave()
                    st.session_state._show_recovery_dialog = False
                    if ok:
                        st.rerun()
            with col2:
                if st.button("✗ Cancel", key="cancel_recovery"):
                    st.session_state._show_recovery_dialog = False
                    st.rerun()

    st.sidebar.markdown("---")
    st.sidebar.markdown("**Import / Export**")

    # Import section
    st.sidebar.markdown("*Import text:*")
    uploaded_file = st.sidebar.file_uploader(
        "Choose file",
        type=["txt", "md", "docx", "pdf"],
        key="file_uploader",
        label_visibility="collapsed",
    )

    if uploaded_file is not None:
        file_ext = uploaded_file.name.split(".")[-1].lower()

        # Show confirmation dialog if current work exists
        current_text = st.session_state.get("main_text", "")
        if current_text and current_text.strip():
            if not st.session_state.get("_confirm_import"):
                st.sidebar.warning("⚠️ This will replace your current text!")
                col1, col2 = st.sidebar.columns(2)
                with col1:
                    if st.button("✓ Import", key="confirm_import_btn"):
                        st.session_state._confirm_import = True
                        st.rerun()
                with col2:
                    if st.button("✗ Cancel", key="cancel_import_btn"):
                        st.session_state._confirm_import = False
                        st.rerun()
        else:
            st.session_state._confirm_import = True

        # Proceed with import if confirmed
        if st.session_state.get("_confirm_import"):
            imported_text = None

            if file_ext == "txt":
                imported_text = import_text_from_txt(uploaded_file)
            elif file_ext == "md":
                imported_text = import_text_from_md(uploaded_file)
            elif file_ext == "docx":
                imported_text = import_text_from_docx(uploaded_file)
            elif file_ext == "pdf":
                imported_text = import_text_from_pdf(uploaded_file)

            if imported_text:
                st.session_state.main_text = imported_text
                # Reset undo history after import
                st.session_state._undo_history = [imported_text]
                st.session_state._redo_history = []
                st.session_state._confirm_import = False
                mark_dirty()
                st.sidebar.success(f"✓ Imported {uploaded_file.name}")
                st.rerun()

    # Export section
    st.sidebar.markdown("*Export text:*")
    export_cols = st.sidebar.columns(4)

    current_text = st.session_state.get("main_text", "")
    export_disabled = not current_text or not current_text.strip()

    # Generate base filename
    title = st.session_state.get("workspace_title", "")
    base_filename = title if title else "olivetti_export"
    # Clean filename
    base_filename = re.sub(r"[^\w\s-]", "", base_filename).strip().replace(" ", "_")

    with export_cols[0]:
        if current_text and not export_disabled:
            txt_data = export_to_txt(current_text, base_filename)
            st.download_button(
                label="TXT",
                data=txt_data,
                file_name=f"{base_filename}.txt",
                mime="text/plain",
                disabled=export_disabled,
                key="export_txt",
            )
        else:
            st.button("TXT", disabled=True, key="export_txt_disabled")

    with export_cols[1]:
        if current_text and not export_disabled:
            md_data = export_to_md(current_text, base_filename)
            st.download_button(
                label="MD",
                data=md_data,
                file_name=f"{base_filename}.md",
                mime="text/markdown",
                disabled=export_disabled,
                key="export_md",
            )
        else:
            st.button("MD", disabled=True, key="export_md_disabled")

    with export_cols[2]:
        if current_text and not export_disabled:
            docx_data = export_to_docx(current_text, base_filename)
            if docx_data:
                st.download_button(
                    label="DOCX",
                    data=docx_data,
                    file_name=f"{base_filename}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    disabled=export_disabled,
                    key="export_docx",
                )
            else:
                st.button("DOCX", disabled=True, key="export_docx_disabled")
        else:
            st.button("DOCX", disabled=True, key="export_docx_disabled2")

    with export_cols[3]:
        if current_text and not export_disabled:
            pdf_data = export_to_pdf(current_text, base_filename)
            if pdf_data:
                st.download_button(
                    label="PDF",
                    data=pdf_data,
                    file_name=f"{base_filename}.pdf",
                    mime="application/pdf",
                    disabled=export_disabled,
                    key="export_pdf",
                )
            else:
                st.button("PDF", disabled=True, key="export_pdf_disabled")
        else:
            st.button("PDF", disabled=True, key="export_pdf_disabled2")

    st.sidebar.markdown("---")
    st.sidebar.markdown("**Cloud Sync (S3)**")

    # Check if cloud sync is configured
    cloud_configured = has_cloud_sync()

    if not cloud_configured:
        st.sidebar.info(
            "☁️ Cloud sync not configured\n\n"
            "Set environment variables or Streamlit secrets:\n"
            "• AWS_ACCESS_KEY_ID\n"
            "• AWS_SECRET_ACCESS_KEY\n"
            "• AWS_S3_BUCKET\n"
            "• AWS_REGION (optional)\n"
            "• AWS_ENDPOINT_URL (optional)"
        )
    else:
        # Cloud sync controls
        cloud_col1, cloud_col2 = st.sidebar.columns(2)

        with cloud_col1:
            if st.button("☁️ Upload", key="cloud_upload_btn"):
                st.session_state._show_cloud_upload = True

        with cloud_col2:
            if st.button("📥 Download", key="cloud_download_btn"):
                st.session_state._show_cloud_download = True

        # Upload dialog
        if st.session_state.get("_show_cloud_upload"):
            with st.sidebar.expander("☁️ Upload to Cloud", expanded=True):
                save_name = st.text_input(
                    "Save name",
                    value=st.session_state.get(
                        "workspace_title",
                        f"save_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                    ),
                    key="cloud_save_name_input",
                )

                # Sanitize save name
                save_name_clean = (
                    re.sub(r"[^\w\s-]", "", save_name).strip().replace(" ", "_")
                )

                if save_name_clean:
                    st.caption(f"Will save as: {save_name_clean}.json")

                col1, col2 = st.columns(2)
                with col1:
                    if st.button(
                        "✓ Upload",
                        key="confirm_cloud_upload",
                        disabled=not save_name_clean,
                    ):
                        if upload_to_cloud(save_name_clean):
                            st.success(f"✓ Uploaded to cloud: {save_name_clean}")
                            st.session_state._show_cloud_upload = False
                            time.sleep(1)
                            st.rerun()
                with col2:
                    if st.button("✗ Cancel", key="cancel_cloud_upload"):
                        st.session_state._show_cloud_upload = False
                        st.rerun()

        # Download dialog
        if st.session_state.get("_show_cloud_download"):
            with st.sidebar.expander("📥 Download from Cloud", expanded=True):
                # List available saves
                cloud_saves = list_cloud_saves()

                if not cloud_saves:
                    st.info("No cloud saves found")
                else:
                    st.markdown("**Available saves:**")

                    # Display saves
                    for save in cloud_saves:
                        col1, col2 = st.columns([3, 1])
                        with col1:
                            if st.button(
                                f"📄 {save['name']}",
                                key=f"select_{save['key']}",
                                help=f"Modified: {save['modified']}\nSize: {save['size']} bytes",
                            ):
                                st.session_state._selected_cloud_save = save
                        with col2:
                            if st.button(
                                "🗑️", key=f"delete_{save['key']}", help="Delete"
                            ):
                                if delete_from_cloud(save["key"]):
                                    st.success("Deleted")
                                    time.sleep(0.5)
                                    st.rerun()

                # Confirmation for selected save
                selected = st.session_state.get("_selected_cloud_save")
                if selected:
                    st.warning(
                        f"⚠️ Download '{selected['name']}'?\nThis will replace your current work!"
                    )
                    st.caption(f"Modified: {selected['modified']}")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("✓ Download", key="confirm_cloud_download"):
                            if download_from_cloud(selected["key"]):
                                st.success(f"✓ Downloaded: {selected['name']}")
                                st.session_state._selected_cloud_save = None
                                st.session_state._show_cloud_download = False
                                time.sleep(1)
                                st.rerun()
                    with col2:
                        if st.button("✗ Cancel", key="cancel_cloud_download_confirm"):
                            st.session_state._selected_cloud_save = None
                            st.rerun()

                st.markdown("---")
                if st.button("Close", key="close_cloud_download"):
                    st.session_state._show_cloud_download = False
                    st.session_state._selected_cloud_save = None
                    st.rerun()

    st.sidebar.markdown("---")
    st.sidebar.markdown("**Voice controls**")
    st.sidebar.checkbox(
        "Enable writing style",
        value=st.session_state.vb_style_on,
        key="vb_style_on",
        on_change=mark_dirty,
    )

    # Check contract for locks
    contract_status = get_contract_status()
    style_locked = contract_status["style_locked"] or contract_status["locked"]
    voice_locked = contract_status["voice_locked"] or contract_status["locked"]
    intensity_limits = contract_status["intensity_limits"]

    writing_styles = ["Neutral", "Crisp", "Flowing"]
    current_style = st.session_state.get("writing_style", "Neutral")
    style_idx = (
        writing_styles.index(current_style) if current_style in writing_styles else 0
    )
    st.sidebar.selectbox(
        "Writing style",
        writing_styles,
        index=style_idx,
        key="writing_style",
        on_change=mark_dirty,
        disabled=style_locked,
        help="🔒 Locked by System Contract" if style_locked else None,
    )
    st.sidebar.slider(
        "AI intensity",
        min_value=intensity_limits["min"],
        max_value=intensity_limits["max"],
        value=float(st.session_state.ai_intensity),
        step=0.01,
        key="ai_intensity",
        on_change=mark_dirty,
        help=(
            f"Range limited to {intensity_limits['min']:.2f}-{intensity_limits['max']:.2f}"
            if intensity_limits["min"] > 0.0 or intensity_limits["max"] < 1.0
            else None
        ),
    )

    st.sidebar.markdown("---")
    st.sidebar.checkbox(
        "Lock story bible editing",
        value=bool(st.session_state.story_bible_lock),
        key="story_bible_lock",
        on_change=mark_dirty,
    )

    # System Contract Panel
    st.sidebar.markdown("---")
    st.sidebar.markdown("**⚖️ System Contract**")

    contract_status = get_contract_status()

    # Show contract status
    if contract_status["locked"]:
        st.sidebar.caption("🔒 Project settings locked")
    else:
        st.sidebar.caption("🔓 Project settings unlocked")

    # Toggle contract panel
    if st.sidebar.button(
        "⚙️ Manage Contract", help="View and manage project integrity settings"
    ):
        st.session_state._show_contract_panel = not st.session_state.get(
            "_show_contract_panel", False
        )

    # Contract management panel
    if st.session_state.get("_show_contract_panel"):
        with st.sidebar.expander("System Contract Settings", expanded=True):
            st.caption("**Project Integrity Controls**")

            # Master lock/unlock
            if contract_status["locked"]:
                if st.button("🔓 Unlock All Settings", key="unlock_all_settings"):
                    unlock_project_settings()
                    st.rerun()
            else:
                if st.button("🔒 Lock All Settings", key="lock_all_settings"):
                    lock_project_settings()
                    st.rerun()

            st.markdown("---")
            st.caption("**AI Intensity Limits**")
            col1, col2 = st.columns(2)
            with col1:
                min_intensity = st.number_input(
                    "Min",
                    min_value=0.0,
                    max_value=1.0,
                    value=contract_status["intensity_limits"]["min"],
                    step=0.05,
                    key="contract_min_intensity",
                )
            with col2:
                max_intensity = st.number_input(
                    "Max",
                    min_value=0.0,
                    max_value=1.0,
                    value=contract_status["intensity_limits"]["max"],
                    step=0.05,
                    key="contract_max_intensity",
                )

            if st.button("Apply Limits", key="apply_intensity_limits"):
                set_intensity_limits(min_intensity, max_intensity)
                st.success("Intensity limits applied")
                st.rerun()

            st.markdown("---")
            st.caption("**Individual Locks**")

            # Style lock
            if contract_status["style_locked"]:
                if st.button("🔓 Unlock Writing Style", key="unlock_style"):
                    unlock_writing_style()
                    st.rerun()
            else:
                if st.button("🔒 Lock Writing Style", key="lock_style"):
                    lock_writing_style()
                    st.rerun()

            # Voice lock
            if contract_status["voice_locked"]:
                if st.button("🔓 Unlock Voice Selection", key="unlock_voice"):
                    unlock_voice_selection()
                    st.rerun()
            else:
                if st.button("🔒 Lock Voice Selection", key="lock_voice"):
                    lock_voice_selection()
                    st.rerun()

            st.markdown("---")
            if st.button(
                "🔄 Reset Contract",
                key="reset_contract",
                help="Reset to default settings",
            ):
                reset_system_contract()
                st.success("Contract reset to defaults")
                st.rerun()

            if st.button("✗ Close", key="close_contract_panel"):
                st.session_state._show_contract_panel = False
                st.rerun()

    # Top bar
    col1, col2 = st.columns([3, 1])
    with col1:
        st.header(st.session_state.get("project_title", "Olivetti Desk"))
    with col2:
        if st.button("Autosave now"):
            autosave_state()
            st.success("Autosaved")

    # Main layout
    left, center, right = st.columns([1.2, 2.4, 1.2])

    with left:
        st.subheader("Story bible")
        st.text_input(
            "Workspace title",
            value=st.session_state.workspace_title,
            key="workspace_title",
            on_change=mark_dirty,
        )
        locked = bool(st.session_state.story_bible_lock)
        st.text_area(
            "Synopsis",
            value=st.session_state.synopsis,
            height=110,
            key="synopsis",
            disabled=locked,
            on_change=mark_dirty,
        )
        st.text_area(
            "Genre/style notes",
            value=st.session_state.genre_style_notes,
            height=110,
            key="genre_style_notes",
            disabled=locked,
            on_change=mark_dirty,
        )
        st.text_area(
            "World",
            value=st.session_state.world,
            height=110,
            key="world",
            disabled=locked,
            on_change=mark_dirty,
        )
        st.text_area(
            "Characters",
            value=st.session_state.characters,
            height=110,
            key="characters",
            disabled=locked,
            on_change=mark_dirty,
        )
        st.text_area(
            "Outline",
            value=st.session_state.outline,
            height=110,
            key="outline",
            disabled=locked,
            on_change=mark_dirty,
        )
        a, b = st.columns(2)
        with a:
            if st.button("Reset workspace"):
                reset_workspace_story_bible()
                st.rerun()
        with b:
            if st.button("Save workspace"):
                save_workspace_from_session()
                autosave_state()
                st.success("Workspace saved")

    with center:
        st.subheader("Writing desk")

        # Bay indicator with visual styling
        current_bay = st.session_state.get("active_bay", "NEW")
        bay_colors = {
            "NEW": "#4A90E2",  # Blue
            "ROUGH": "#F5A623",  # Orange
            "EDIT": "#7ED321",  # Green
            "FINAL": "#BD10E0",  # Purple
        }
        bay_color = bay_colors.get(current_bay, "#888888")
        st.markdown(
            f"""
            <div style="
                background-color: {bay_color};
                color: white;
                padding: 8px 16px;
                border-radius: 8px;
                text-align: center;
                font-weight: bold;
                font-size: 14px;
                margin-bottom: 12px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            ">
                📍 Current Bay: {current_bay}
            </div>
            """,
            unsafe_allow_html=True,
        )

        # Word count indicators
        current_text = st.session_state.get("main_text", "")
        current_word_count = count_words(current_text)

        # Get word counts for all bays
        bay_word_counts = {}
        for bay in ["ROUGH", "EDIT", "FINAL"]:
        workspace = st.session_state.active_project_by_bay.get(bay)
            if workspace:
                bay_text = workspace.get("main_text", "")
                bay_word_counts[bay] = count_words(bay_text)
            else:
                bay_word_counts[bay] = 0

        # Display word counts and trash bin indicator
        wc_col1, wc_col2, wc_col3 = st.columns([1, 1, 0.5])
        with wc_col1:
            st.markdown(
                f"**Current Draft:** {current_word_count:,} words",
                help="Live word count for the active draft",
            )
        with wc_col2:
            bay_counts_text = " • ".join(
                [f"{bay}: {count:,}" for bay, count in bay_word_counts.items()]
            )
            st.markdown(
                f"**Bay Counts:** {bay_counts_text}",
                help="Word counts for drafts in each bay",
            )
        with wc_col3:
            # Trash bin indicator
            trash_items = get_trash_bin_items()
            trash_count = len(trash_items)
            if trash_count > 0:
                trash_color = "#E74C3C"  # Red for items in trash
                st.markdown(
                    f"""
                    <div style="
                        background-color: {trash_color};
                        color: white;
                        padding: 6px 12px;
                        border-radius: 6px;
                        text-align: center;
                        font-weight: bold;
                        font-size: 13px;
                        cursor: pointer;
                        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                    ">
                        🗑️ {trash_count}
                    </div>
                    """,
                    unsafe_allow_html=True,
                    help=f"{trash_count} item(s) in Trash Bin",
                )

        st.markdown("---")

        # Top bar: Undo/Redo + Bay Transfer + Export + Delete
        top_col1, top_col2, top_col3, top_col4, top_col5 = st.columns([1, 1, 2, 2, 1])

        # Undo/Redo controls
        with top_col1:
            if st.button("↶ Undo", disabled=not can_undo(), help="Undo (Ctrl+Z)"):
                undo_text()
                st.rerun()
        with top_col2:
            if st.button("↷ Redo", disabled=not can_redo(), help="Redo (Ctrl+Y)"):
                redo_text()
                st.rerun()

        # Bay Transfer button
        with top_col3:
            current_bay = st.session_state.get("active_bay", "NEW")
            next_bay = get_next_bay()
            if next_bay:
                transfer_label = f"{current_bay} → {next_bay}"
                if st.button(transfer_label, help=f"Move draft to {next_bay} bay"):
                    if not transfer_to_next_bay():
                        # Need confirmation - will show dialog below
                        st.rerun()
                    else:
                        st.success(f"✓ Moved to {next_bay} bay")
                        st.rerun()
            else:
                st.button("Transfer", disabled=True, help="No next bay available")

        # Export button
        with top_col4:
            has_text = bool(st.session_state.get("main_text", "").strip())
            if has_text:
                export_filename = get_export_filename()
                if st.button("📤 Export", help=f"Export as {export_filename}"):
                    st.session_state._show_export_dialog = True
            else:
                st.button("📤 Export", disabled=True, help="No text to export")

        # Delete button
        with top_col5:
            has_text = bool(st.session_state.get("main_text", "").strip())
            if has_text:
                if st.button("🗑️ Delete", help="Delete current draft"):
                    st.session_state._confirm_delete = True
            else:
                st.button("🗑️ Delete", disabled=True, help="No draft to delete")

        # Transfer confirmation dialog
        if st.session_state.get("_transfer_needs_confirmation"):
            target_bay = st.session_state.get("_transfer_target_bay")
            with st.expander(f"⚠️ {target_bay} Bay Already Has Content", expanded=True):
                st.warning(
                    f"The {target_bay} bay already contains a draft. Transfer will overwrite it."
                )
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("✓ Overwrite", key="confirm_transfer"):
                        _execute_bay_transfer(target_bay)
                        st.session_state._transfer_needs_confirmation = False
                        st.session_state._transfer_target_bay = None
                        st.success(f"✓ Moved to {target_bay} bay")
                        st.rerun()
                with col2:
                    if st.button("✗ Cancel", key="cancel_transfer"):
                        st.session_state._transfer_needs_confirmation = False
                        st.session_state._transfer_target_bay = None
                        st.rerun()

        # Delete confirmation dialog
        if st.session_state.get("_confirm_delete"):
            current_bay = st.session_state.get("active_bay", "NEW")
            word_count = count_words(st.session_state.get("main_text", ""))
            with st.expander(f"⚠️ Delete Draft from {current_bay} Bay?", expanded=True):
                st.info(
                    f"Draft will be moved to Trash Bin ({word_count:,} words). You can restore it later."
                )
                col1, col2 = st.columns(2)
                with col1:
                    if st.button(
                        "✓ Move to Trash", key="confirm_delete_btn", type="primary"
                    ):
                        if delete_current_draft():
                            st.session_state._confirm_delete = False
                            st.success(f"✓ Draft moved to Trash Bin")
                            st.rerun()
                        else:
                            st.session_state._confirm_delete = False
                            st.rerun()
                with col2:
                    if st.button("✗ Cancel", key="cancel_delete_btn"):
                        st.session_state._confirm_delete = False
                        st.rerun()

        # Trash Bin notification and viewer
        trash_items = get_trash_bin_items()
        if trash_items:
            # Trash bin toggle button
            trash_button_col, _ = st.columns([1, 5])
            with trash_button_col:
                if st.button(
                    f"🗑️ Trash Bin ({len(trash_items)})",
                    help="View deleted drafts",
                    key="toggle_trash_bin",
                ):
                    st.session_state._show_trash_bin = not st.session_state.get(
                        "_show_trash_bin", False
                    )

            # Trash bin viewer
            if st.session_state.get("_show_trash_bin"):
                with st.expander("🗑️ Trash Bin Contents", expanded=True):
                    st.caption(
                        f"Showing {len(trash_items)} deleted draft(s). Maximum 10 items kept."
                    )

                    for item in trash_items:
                        bay = item.get("bay", "unknown")
                        workspace = item.get("workspace", {})
                        title = workspace.get("workspace_title", "Untitled")
                        text = workspace.get("main_text", "")
                        word_count = count_words(text)
                        timestamp = item.get("timestamp", "")
                        trash_id = item.get("id", "")

                        # Format timestamp
                        try:
                            dt = datetime.fromisoformat(timestamp)
                            time_str = dt.strftime("%b %d, %Y %I:%M %p")
                        except:
                            time_str = "Unknown date"

                        st.markdown("---")
                        st.markdown(f"**{title}** ({bay} bay)")
                        st.caption(f"{word_count:,} words • Deleted: {time_str}")

                        # Preview text (first 100 chars)
                        preview = text[:100] + "..." if len(text) > 100 else text
                        if preview:
                            st.text(preview)

                        # Action buttons
                        col1, col2 = st.columns(2)
                        with col1:
                            if st.button(
                                "↶ Restore",
                                key=f"restore_{trash_id}",
                                help=f"Restore to {bay} bay",
                            ):
                                if restore_from_trash(trash_id):
                                    st.success(f"✓ Restored to {bay} bay")
                                    st.rerun()
                        with col2:
                            if st.button(
                                "🗑️ Delete Permanently",
                                key=f"delete_perm_{trash_id}",
                                help="Permanently delete this draft",
                            ):
                                st.session_state._confirm_permanent_delete = trash_id

                    # Clear all trash button
                    st.markdown("---")
                    if st.button(
                        "🗑️ Empty Trash Bin", key="empty_trash_bin", type="secondary"
                    ):
                        clear_trash_bin()
                        st.session_state._show_trash_bin = False
                        st.success("✓ Trash Bin emptied")
                        st.rerun()

        # Permanent delete confirmation
        if st.session_state.get("_confirm_permanent_delete"):
            trash_id = st.session_state._confirm_permanent_delete
            with st.expander("⚠️ Permanently Delete Draft?", expanded=True):
                st.error(
                    "This will permanently delete this draft. This action cannot be undone."
                )
                col1, col2 = st.columns(2)
                with col1:
                    if st.button(
                        "✓ Delete Permanently",
                        key="confirm_perm_delete",
                        type="primary",
                    ):
                        if permanently_delete_from_trash(trash_id):
                            st.session_state._confirm_permanent_delete = None
                            st.success("✓ Draft permanently deleted")
                            st.rerun()
                with col2:
                    if st.button("✗ Cancel", key="cancel_perm_delete"):
                        st.session_state._confirm_permanent_delete = None
                        st.rerun()

        # Export dialog
        if st.session_state.get("_show_export_dialog"):
            with st.expander("📤 Export Draft", expanded=True):
                export_filename = get_export_filename()
                st.markdown(f"**Export as:** {export_filename}")
                st.caption(f"Current bay: {st.session_state.get('active_bay', 'NEW')}")

                # Export format selection
                export_format = st.radio(
                    "Format",
                    ["TXT", "MD", "DOCX", "PDF"],
                    horizontal=True,
                    key="export_format_choice",
                )

                # Generate export data
                current_text = st.session_state.get("main_text", "")
                export_data = None
                mime_type = "text/plain"
                file_ext = export_format.lower()

                if export_format == "TXT":
                    export_data = export_to_txt(current_text, export_filename)
                    mime_type = "text/plain"
                elif export_format == "MD":
                    export_data = export_to_md(current_text, export_filename)
                    mime_type = "text/markdown"
                elif export_format == "DOCX":
                    export_data = export_to_docx(current_text, export_filename)
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                elif export_format == "PDF":
                    export_data = export_to_pdf(current_text, export_filename)
                    mime_type = "application/pdf"

                col1, col2 = st.columns(2)
                with col1:
                    if export_data:
                        st.download_button(
                            label=f"⬇ Download {export_format}",
                            data=export_data,
                            file_name=f"{export_filename}.{file_ext}",
                            mime=mime_type,
                            key="export_download_btn",
                        )
                    else:
                        st.button(f"⬇ Download {export_format}", disabled=True)
                with col2:
                    if st.button("Close", key="close_export_dialog"):
                        st.session_state._show_export_dialog = False
                        st.rerun()

        st.text_area(
            "Main text",
            value=st.session_state.main_text,
            height=440,
            key="main_text",
            on_change=on_main_text_change,
        )
        st.markdown("**Actions**")
        action_col1, action_col2, action_col3 = st.columns(3)
        if action_col1.button("WRITE"):
            do_action("WRITE")
        if action_col2.button("REWRITE"):
            do_action("REWRITE")
        if action_col3.button("EXPAND"):
            do_action("EXPAND")
        st.markdown("**Tool output**")
        st.text_area(
            "Tool output",
            value=st.session_state.tool_output,
            height=160,
            key="tool_output_display",
        )

    with right:
        st.subheader("Voice bible")

        # Check contract for voice lock
        contract_status = get_contract_status()
        voice_locked = contract_status["voice_locked"] or contract_status["locked"]

        # Voice selection
        names = voice_names_for_selector()
        idx = 0
        if st.session_state.get("trained_voice") in names:
            idx = names.index(st.session_state.get("trained_voice"))
        st.selectbox(
            "Trained voice",
            options=names,
            index=idx,
            key="trained_voice",
            on_change=mark_dirty,
            disabled=voice_locked,
            help="🔒 Locked by System Contract" if voice_locked else None,
        )
        current_lane = st.session_state.get("voice_lane", "Narration")
        lane_idx = (
            LANES.index(current_lane)
            if current_lane in LANES
            else LANES.index("Narration")
        )
        st.selectbox(
            "Lane",
            options=LANES,
            index=lane_idx,
            key="voice_lane",
            on_change=mark_dirty,
        )
        st.text_area(
            "Voice sample",
            value=st.session_state.voice_sample,
            height=140,
            key="voice_sample",
            on_change=mark_dirty,
        )
        if st.button("Add voice sample"):
            ok = add_voice_sample(
                st.session_state.trained_voice,
                st.session_state.voice_lane,
                st.session_state.voice_sample,
            )
            if ok:
                st.success("Sample added")
            else:
                st.warning("Pick a voice (not None) and add non-empty text.")

        # My Voice Profile Management
        st.markdown("---")
        st.caption("**My Voice Profiles**")
        my_voices = get_my_voice_profiles()

        if my_voices:
            st.caption(f"{len(my_voices)} voice profile(s)")

            # Show selected voice stats
            current_voice = st.session_state.get("trained_voice", "— None —")
            if current_voice in my_voices:
                sample_count = get_voice_sample_count(current_voice)
                voice_stats = get_voice_learning_stats(current_voice)
                st.caption(f"• {sample_count} training sample(s)")
                st.caption(f"• {voice_stats.get('total_edits', 0)} edit(s) learned")

        # Create new My Voice profile
        col1, col2 = st.columns([2, 1])
        with col1:
            new_voice_name = st.text_input(
                "New voice name",
                key="new_voice_name_input",
                placeholder="My Writing Style",
                label_visibility="collapsed",
            )
        with col2:
            if st.button("➕ Create", help="Create new My Voice profile"):
                if new_voice_name and new_voice_name.strip():
                    if create_my_voice_profile(new_voice_name.strip()):
                        st.success(f"Created: {new_voice_name}")
                        st.session_state.new_voice_name_input = ""
                        st.rerun()
                    else:
                        st.error("Voice already exists")

        # Delete current voice if it's a My Voice profile
        current_voice = st.session_state.get("trained_voice", "— None —")
        if current_voice in my_voices:
            if st.button("🗑️ Delete Voice", help=f"Delete {current_voice} profile"):
                st.session_state._confirm_delete_voice = current_voice

        # Delete voice confirmation
        if st.session_state.get("_confirm_delete_voice"):
            voice_to_delete = st.session_state._confirm_delete_voice
            with st.expander(f"⚠️ Delete '{voice_to_delete}'?", expanded=True):
                sample_count = get_voice_sample_count(voice_to_delete)
                st.warning(
                    f"This will delete the voice profile and all {sample_count} training samples."
                )
                col1, col2 = st.columns(2)
                with col1:
                    if st.button(
                        "✓ Delete", key="confirm_delete_voice_btn", type="primary"
                    ):
                        if delete_my_voice_profile(voice_to_delete):
                            st.session_state._confirm_delete_voice = None
                            if st.session_state.get("trained_voice") == voice_to_delete:
                                st.session_state.trained_voice = "— None —"
                            st.success(f"Deleted: {voice_to_delete}")
                            st.rerun()
                with col2:
                    if st.button("✗ Cancel", key="cancel_delete_voice_btn"):
                        st.session_state._confirm_delete_voice = None
                        st.rerun()

        st.markdown("---")
        st.caption(
            f"Detected lane: {current_lane_from_draft(st.session_state.main_text)}"
        )

        # Adaptive Style Learning Panel
        st.markdown("---")
        st.subheader("Style Learning")

        learning_stats = get_style_learning_stats()

        # Learning toggle
        learning_enabled = st.checkbox(
            "Enable adaptive learning",
            value=learning_stats["enabled"],
            key="style_learning_toggle",
            help="Learn from your edits, rewrites, and preferences over time",
        )
        if learning_enabled != learning_stats["enabled"]:
            toggle_style_learning(learning_enabled)

        if learning_enabled:
            # Show learning statistics
            st.caption(f"**Learning Stats:**")
            st.caption(f"• Total edits learned: {learning_stats['total_edits']}")
            st.caption(f"• Accepted suggestions: {learning_stats['total_accepts']}")
            st.caption(f"• Rejected suggestions: {learning_stats['total_rejects']}")
            st.caption(f"• Learned words: {learning_stats['learned_words']}")

            if learning_stats["avg_sentence_length"] > 0:
                st.caption(
                    f"• Avg sentence length: {learning_stats['avg_sentence_length']:.1f} words"
                )

            # Learning controls
            col1, col2 = st.columns(2)
            with col1:
                if st.button(
                    "📊 View Patterns",
                    help="View learned style patterns",
                    key="view_patterns_btn",
                ):
                    st.session_state._show_style_patterns = not st.session_state.get(
                        "_show_style_patterns", False
                    )
            with col2:
                if st.button(
                    "🔄 Reset", help="Reset learned styles", key="reset_learning_btn"
                ):
                    st.session_state._confirm_reset_learning = True

            # Show patterns if toggled
            if st.session_state.get("_show_style_patterns"):
                with st.expander("📊 Learned Patterns", expanded=True):
                    engine = st.session_state.get("_style_learning_engine")
                    if engine:
                        patterns = engine.get("learned_patterns", {})

                        # Sentence length preference
                        sent_prefs = patterns.get("sentence_length", {})
                        if any(sent_prefs.values()):
                            st.caption("**Sentence Length:**")
                            total = sum(sent_prefs.values())
                            if total > 0:
                                st.caption(
                                    f"Short: {sent_prefs.get('short', 0)/total*100:.0f}% • "
                                    f"Medium: {sent_prefs.get('medium', 0)/total*100:.0f}% • "
                                    f"Long: {sent_prefs.get('long', 0)/total*100:.0f}%"
                                )

                        # Top words
                        word_prefs = patterns.get("word_preferences", {})
                        if word_prefs:
                            top_words = sorted(
                                word_prefs.items(), key=lambda x: x[1], reverse=True
                            )[:10]
                            st.caption("**Top Words:**")
                            st.caption(", ".join([f"{w} ({c})" for w, c in top_words]))

                    if st.button("Close Patterns", key="close_patterns_btn"):
                        st.session_state._show_style_patterns = False
                        st.rerun()

            # Reset confirmation
            if st.session_state.get("_confirm_reset_learning"):
                with st.expander("⚠️ Reset Style Learning?", expanded=True):
                    st.warning(
                        "This will erase all learned patterns and start fresh. This action cannot be undone."
                    )
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button(
                            "✓ Reset", key="confirm_reset_learning_btn", type="primary"
                        ):
                            reset_style_learning()
                            st.session_state._confirm_reset_learning = False
                            st.success("Style learning reset")
                            st.rerun()
                    with col2:
                        if st.button("✗ Cancel", key="cancel_reset_learning_btn"):
                            st.session_state._confirm_reset_learning = False
                            st.rerun()
        else:
            st.caption("Learning is disabled. Enable to start adapting to your style.")

        st.markdown("---")
        st.write(
            f"Autosave: {st.session_state.get('autosave_time', 'never')} • "
            f"Dirty: {bool(st.session_state.get('_dirty'))} • "
            f"Last action: {st.session_state.get('last_action', '—')}"
        )


# ============================================================
# STARTUP
# ============================================================
def main() -> None:
    init_state()

    # Automatic recovery prompt on first run
    if not st.session_state.get("_autosave_checked"):
        st.session_state._autosave_checked = True
        autosave_info = get_autosave_info()
        # Only prompt if autosave exists and has content
        if autosave_info and autosave_info.get("has_content"):
            # Check if current session is empty (fresh start)
            current_text = st.session_state.get("main_text", "")
            if not current_text or not current_text.strip():
                # Auto-load on fresh start
                if load_autosave():
                    st.toast("✓ Recovered autosave from " + autosave_info["saved_ts"])
            else:
                # Show dialog if there's existing work
                st.session_state._show_recovery_dialog = True

    # Keep workspace mirror consistent
    save_workspace_from_session()
    # Throttled background autosave for edits
    maybe_autosave_throttled()
    main_ui()


if __name__ == "__main__":
    main()
